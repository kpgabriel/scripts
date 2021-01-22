import pandas as pd
import re
from xlsxwriter.utility import xl_rowcol_to_cell
from docx import Document
from docx.shared import Inches




try:
    xls = pd.ExcelFile('C:\Projects\Reports\\2021\Q1\\10-K\Financial_Report_Guskin Gold Corp..xlsx')
    form_pattern = re.compile('(Consolidated|Statements)')
    sheets = xls.sheet_names
    newlist = list(filter(form_pattern.match, sheets))
    newlist.insert(0,sheets[0])
    
    dict_sheets = pd.read_excel('C:\Projects\Reports\\2021\Q1\\10-K\Financial_Report_Guskin Gold Corp..xlsx',sheet_name=newlist)
    
    # print(newlist)

    income_sheets = {}
    other_sheets_dict = {}
    for x in range(len(newlist)):
        # Create the dictionary of data frame objects with sheet names
        income_sheets['{0}'.format(newlist[x])] =  pd.DataFrame(dict_sheets[newlist[x]])

    # result of all other sheets is a list, order is preserved
    other_sheets = [i for i in sheets if i not in income_sheets] 
    other_dict_sheets = pd.read_excel('C:\Projects\Reports\\2021\Q1\\10-K\Financial_Report_Guskin Gold Corp..xlsx', sheet_name=other_sheets)
    for x in range(len(other_sheets)):
        other_sheets_dict['{0}'.format(other_sheets[x])] =  pd.DataFrame(other_dict_sheets[other_sheets[x]])

    # print(income_sheets.keys())
    # print(other_sheets_dict.keys())
    # Make a writer using xlsxwriter library
    writer_orig = pd.ExcelWriter('C:/Projects/Reports/simple.xlsx', engine='xlsxwriter')

    # Simple excel document
    for sheet_name in income_sheets.keys():
        income_sheets[sheet_name].to_excel(writer_orig, sheet_name=sheet_name, index=False)
    
    writer_orig.save()

    # Now we format and write a more complicated output

    writer = pd.ExcelWriter('C:/Projects/Reports/fancy.xlsx', engine='xlsxwriter')
    # Get the workbook
    workbook = writer.book

    # Add a bold format to use to highlight cells.
    bold = workbook.add_format({'bold': True})
    # Add a custom format for cells with money.
    money = workbook.add_format({'num_format': '_("$ "#,##0_);_("$ "(#,##0)'})

    num_sheet = 0
    for sheet_name in income_sheets.keys():
        income_sheets[sheet_name].to_excel(writer, sheet_name=sheet_name, index=False)

        # Get the worksheet
        worksheet = writer.sheets[sheet_name]


        if sheet_name == 'Document And Entity Information':
            cell = xl_rowcol_to_cell(7, 1)   # B8
            
        # Change the cell columns in all sheets
        
        worksheet.set_column('A:A', 40)
        worksheet.set_column('B:B', 25)
        worksheet.set_column('C:C', 25)
        worksheet.set_column('D:D', 15)
        worksheet.set_column('E:E', 15)
        

    writer.save()

    # Now we try writing to word 
    document = Document()
    for i in range(len(other_sheets)): 
        document.add_heading('{0}'.format(other_sheets[i]), 0)

        p = document.add_paragraph('{0}'.format(other_sheets_dict[other_sheets[i]]))
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )

    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

        document.add_page_break()

    document.save('C:/Projects/Reports/demo.docx')

except Exception as e:
    print(e)
